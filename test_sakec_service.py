import requests
import base64
import os
from docx import Document
import time

def create_dummy_docx(filename="dummy_test.docx"):
    doc = Document()
    doc.add_heading('Test Document for Case Study', 0)
    doc.add_paragraph('This is a test document containing some engineering context.')
    doc.add_paragraph('The system at the fictional Acme Corp failed due to a lack of redundancy in the network architecture.')
    doc.add_paragraph('Engineers observed packet loss and high latency during peak hours.')
    doc.save(filename)
    return filename

def test_generate_case_study():
    filename = create_dummy_docx()
    
    with open(filename, "rb") as f:
        file_bytes = f.read()
        b64_data = base64.b64encode(file_bytes).decode('utf-8')
    
    url = "http://localhost:5016/generate-sakec-case-study"
    payload = {
        "filename": filename,
        "file_data": b64_data,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    
    print(f"Sending request to {url}...")
    try:
        start_time = time.time()
        response = requests.post(url, json=payload)
        end_time = time.time()
        
        print(f"Response Status Code: {response.status_code}")
        print(f"Time Taken: {end_time - start_time:.2f}s")
        
        if response.status_code == 200:
            data = response.json()
            if "completed_file_data" in data:
                print("SUCCESS: Received completed file data.")
                output_filename = data.get("completed_filename", "output_case_study.docx")
                
                # Decode and save
                with open(output_filename, "wb") as f_out:
                    f_out.write(base64.b64decode(data["completed_file_data"]))
                
                print(f"Saved generated case study to: {output_filename}")
                return True
            else:
                print("FAILURE: Response missing 'completed_file_data'.")
                print(data)
                return False
        else:
            print("FAILURE: Request failed.")
            print(response.text)
            return False
            
    except Exception as e:
        print(f"EXCEPTION: {e}")
        return False
    finally:
        if os.path.exists(filename):
            os.remove(filename)

if __name__ == "__main__":
    test_generate_case_study()
