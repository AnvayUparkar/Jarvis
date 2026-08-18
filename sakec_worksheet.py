# sakec_worksheet.py
import os
import uuid
import base64
import json
import PyPDF2
import pytesseract
import random # NEW: For uniqueness entropy
import google.generativeai as genai
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename
from PIL import Image

# --- API Key Configuration ---
try:
    from apikey import GEN_AI_API_KEY
    if not GEN_AI_API_KEY:
        raise ValueError("GEN_AI_API_KEY is empty in apikey.py.")
except ImportError:
    print("WARNING: apikey.py not found. Please set your API key directly.")
    GEN_AI_API_KEY = os.getenv("GEMINI_API_KEY", "")
    if not GEN_AI_API_KEY:
        print("CRITICAL ERROR: GEMINI_API_KEY is not set.")

# --- Tesseract OCR Configuration ---
try:
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' 
    pass
except Exception as e:
    print(f"WARNING: pytesseract configuration failed: {e}")

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "www/uploaded_files"
WORKSHEET_FOLDER = "www/worksheets"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(WORKSHEET_FOLDER, exist_ok=True)

# Configure Gemini AI model
if GEN_AI_API_KEY:
    try:
        genai.configure(api_key=GEN_AI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
    except Exception as e:
        print(f"CRITICAL ERROR: Failed to configure Gemini API: {e}")
        model = None
else:
    print("CRITICAL ERROR: Gemini API key is missing.")
    model = None

# --- Text Extraction Helper Functions (Reused) ---
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"ERROR: Could not extract text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_image(image_path):
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text
    except pytesseract.TesseractNotFoundError:
        print("ERROR: Tesseract OCR not found.")
        return ""
    except Exception as e:
        print(f"ERROR: Could not perform OCR on image {image_path}: {e}")
        return ""

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"ERROR: Could not extract text from PDF {file_path}: {e}")
        return ""

# --- SAKEC/MU Scenario Logic ---
def generate_sakec_questions(extracted_text, mime_type, base64_file_data):
    if model is None:
        raise Exception("Gemini API model is not configured.")

    gemini_input_parts = []
    
    if mime_type.startswith("image/"):
        gemini_input_parts.append({
            "inline_data": {
                "mime_type": mime_type,
                "data": base64_file_data
            }
        })
        gemini_input_parts.append({
            "text": f"Analyze this image content. Extracted text:\n\n{extracted_text}"
        })
    else:
        gemini_input_parts.append({"text": extracted_text})

    prompt = f"""
    You are a senior university professor creating a high-level Scenario-Based Assignment aimed at engineering students (SAKEC / MU standards).
    
    Using ONLY the provided document content, generate 5 to 8 scenario-based questions.
    Do NOT use outside knowledge or hallucinate topics not present in the content.

    ### ENGINEERING SCENARIO FORMATTING RULE (MANDATORY)
    Frame every scenario strictly from an engineering student / engineer’s perspective, following university exam-style narration:

    1. **Start scenarios with realistic academic or industry contexts** such as:
       - "During a computer networks laboratory experiment…"
       - "While analyzing network performance in an organization…"
       - "When an engineering student studies protocol behavior…"
       - "During deployment of a campus / office network…"

    2. **Keep the scenario short, focused, and technical (2–3 lines max).** Avoid storytelling, casual language, or business-only narratives.

    3. **Immediately follow the scenario with a direct, marks-based analytical question.**
       - Questions must be phrased like engineering examination questions:
         - "Explain the working of …"
         - "Compare … with respect to …"
         - "How does … improve …"
         - "Discuss protocol behavior, advantages, limitations, and use-cases"

    ### CRITICAL: UNIQUENESS & ENTROPY REQUIREMENT (95% Deviation)
    Ensure that this generated assignment achieves at least 95% uniqueness by significantly varying scenario context, stakeholders, and failure conditions.

    ### DYNAMIC SCENARIO ANGLE: {random.choice(['Industrial Application', 'Research & Development', 'Startup/Innovation', 'Safety & Failure Analysis', 'Environmental Sustainability', 'Urban Infrastructure', 'Global Supply Chain'])}
    (Force all scenarios to loosely align with this specific angle).

    Return the output as a JSON ARRAY of objects with this structure:
    [
      {{
        "scenario": "During a laboratory experiment on [Topic]...",
        "marks": 7,
        "questions": [
          "a) Explain the working of...",
          "b) Compare [X] with respect to [Y]...",
          "c) Discuss protocol behavior..."
        ],
        "topic": "Specific topic name inferred from content",
        "difficulty": "Medium" or "Hard",
        "suggested_answer": "A structured university-level answer key."
      }}
    ]
    """

    if mime_type.startswith("image/"):
        gemini_input_parts[1]["text"] += "\n\n" + prompt
    else:
        gemini_input_parts.append({"text": prompt})

    try:
        response = model.generate_content(gemini_input_parts)
        text_response = response.text
        
        # Clean markdown
        if text_response.startswith("```json"):
            text_response = text_response[7:].strip()
        if text_response.endswith("```"):
            text_response = text_response[:-3].strip()
            
        return json.loads(text_response)
    except Exception as e:
        raise Exception(f"Gemini generation error: {e}")

def create_sakec_docx(questions_data, filename="sakec_assignment.docx"):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # Title Page
    document.add_heading("Scenario Based Assignment", level=0)
    document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Infer main topic from first question if available
    main_topic = questions_data[0].get('topic', 'General') if questions_data else 'General'
    document.add_paragraph(f"Topic: {main_topic}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Total Questions: {len(questions_data)}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\n")

    # Questions Section
    for i, q in enumerate(questions_data):
        # Scenario Block
        document.add_heading(f"Q{i+1}.", level=1)
        scenario_para = document.add_paragraph()
        scenario_run = scenario_para.add_run(f"Scenario: {q.get('scenario', '')}")
        scenario_run.italic = True
        
        # Question Header
        q_header = document.add_paragraph()
        q_header.add_run(f"\nQuestion ({q.get('marks', 7)} Marks):").bold = True
        
        # Sub-questions
        if isinstance(q.get('questions'), list):
            for sub_q in q['questions']:
                document.add_paragraph(sub_q, style='List Bullet')
        else:
            document.add_paragraph(q.get('questions', ''), style='List Bullet')
            
        document.add_paragraph("\n" + "_"*80 + "\n")

    # Answer Key Section
    document.add_page_break()
    document.add_heading("Answer Key / Marking Scheme", level=0)
    
    for i, q in enumerate(questions_data):
        document.add_heading(f"Q{i+1} Answer:", level=2)
        document.add_paragraph(q.get('suggested_answer', 'No answer provided.'))
        document.add_paragraph("\n")

    file_path = os.path.join(WORKSHEET_FOLDER, filename)
    document.save(file_path)
    return file_path

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "SAKEC Worksheet Service Running", "port": 5015}), 200

@app.route("/generate-sakec-worksheet", methods=["POST"])
def generate_sakec_worksheet():
    if not request.json:
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")
    mime_type = request.json.get("mime_type")

    if not all([filename, base64_file_data, mime_type]):
        return jsonify({"error": "Missing fields"}), 400

    temp_path = None
    docx_path = None
    
    try:
        # Decode and Save Temp File
        file_bytes = base64.b64decode(base64_file_data)
        safe_name = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_path = os.path.join(UPLOAD_FOLDER, safe_name)
        
        with open(temp_path, "wb") as f:
            f.write(file_bytes)

        # Extract Text
        extracted_text = ""
        if mime_type.endswith('document') or filename.endswith('.docx'):
            extracted_text = extract_text_from_docx(temp_path)
        elif mime_type == 'application/pdf' or filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(temp_path)
        elif mime_type.startswith('image/'):
            extracted_text = extract_text_from_image(temp_path)
        
        # Generate Scenarios
        questions_data = generate_sakec_questions(extracted_text, mime_type, base64_file_data)
        
        # Create DOCX
        output_filename = f"SAKEC_Assignment_{uuid.uuid4()}.docx"
        docx_path = create_sakec_docx(questions_data, output_filename)
        
        # Return Result
        with open(docx_path, "rb") as f:
            b64_output = base64.b64encode(f.read()).decode('utf-8')
            
        return jsonify({
            "completed_filename": output_filename,
            "completed_file_data": b64_output
        })

    except Exception as e:
        print(f"ERROR: {e}")
        return jsonify({"error": str(e)}), 500
        
    finally:
        # Cleanup
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
        if docx_path and os.path.exists(docx_path):
            os.remove(docx_path)

if __name__ == "__main__":
    app.run(port=5015, debug=True)
