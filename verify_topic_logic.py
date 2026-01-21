import requests
import base64
from io import BytesIO
from docx import Document

def test_topic_logic(topic="Sustainable Urban Planning"):
    print(f"Testing generation for topic: {topic}")
    
    # Simulate the logic added to main.py
    try:
        # 1. Create Synthetic DOCX
        doc = Document()
        doc.add_heading(f"Topic: {topic}", 0)
        doc.add_paragraph(f"Generate a comprehensive engineering case study centered around the topic: {topic}.")
        doc.add_paragraph("Context: Industrial implementation and challenges.")
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        encoded_content = base64.b64encode(file_stream.read()).decode('utf-8')
        
        # 2. Prepare Payload
        payload = {
            "file_data": encoded_content,
            "filename": f"Topic_{topic}.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        # 3. Send to Service
        url = "http://localhost:5016/generate-sakec-case-study"
        print(f"Sending to {url}...")
        response = requests.post(url, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            if result.get("completed_file_data"):
                filename = result.get("completed_filename", "Topic_Output.docx")
                with open(filename, "wb") as f:
                    f.write(base64.b64decode(result["completed_file_data"]))
                print(f"SUCCESS: Generated case study saved to {filename}")
            else:
                print("FAILURE: No file data in response.")
        else:
            print(f"FAILURE: Service returned {response.status_code}")
            print(response.text)

    except Exception as e:
        print(f"EXCEPTION: {e}")

if __name__ == "__main__":
    test_topic_logic()
