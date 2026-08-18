# sakec_case_study.py
import os
import uuid
import base64
import json
import PyPDF2
import pytesseract
import google.generativeai as genai
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename
from PIL import Image

# --- API Key Configuration ---
try:
    from apikey import GEN_AI_API_KEY
    if not GEN_AI_API_KEY:
        raise ValueError("GEN_AI_API_KEY is empty in apikey.py.")
except ImportError:
    print("WARNING: apikey.py not found. Please set your API key directly.")
    GEN_AI_API_KEY = os.getenv("GEMINI_API_KEY", "")
    if not GEN_AI_API_KEY:
        print("CRITICAL ERROR: GEMINI_API_KEY is not set.")

# --- Tesseract OCR Configuration ---
try:
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' 
    pass
except Exception as e:
    print(f"WARNING: pytesseract configuration failed: {e}")

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "www/uploaded_files"
CASE_STUDY_FOLDER = "www/case_studies"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CASE_STUDY_FOLDER, exist_ok=True)

# Configure Gemini AI model
if GEN_AI_API_KEY:
    try:
        genai.configure(api_key=GEN_AI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash-lite")
    except Exception as e:
        print(f"CRITICAL ERROR: Failed to configure Gemini API: {e}")
        model = None
else:
    print("CRITICAL ERROR: Gemini API key is missing.")
    model = None

# --- Text Extraction Helper Functions ---
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"ERROR: Could not extract text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_image(image_path):
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text
    except pytesseract.TesseractNotFoundError:
        print("ERROR: Tesseract OCR not found.")
        return ""
    except Exception as e:
        print(f"ERROR: Could not perform OCR on image {image_path}: {e}")
        return ""

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"ERROR: Could not extract text from PDF {file_path}: {e}")
        return ""

# --- Case Study Generation Logic ---
def generate_sakec_case_study_content(extracted_text, mime_type, base64_file_data):
    if model is None:
        raise Exception("Gemini API model is not configured.")

    gemini_input_parts = []
    
    if mime_type.startswith("image/"):
        gemini_input_parts.append({
            "inline_data": {
                "mime_type": mime_type,
                "data": base64_file_data
            }
        })
        gemini_input_parts.append({
            "text": f"Analyze this image content. Extracted text:\n\n{extracted_text}"
        })
    else:
        gemini_input_parts.append({"text": extracted_text})

    prompt = """
    You are a senior university professor creating a formal SAKEC / MU-compliant Engineering Case Study based strictly on the provided document.

    OBJECTIVE:
    Generate a detailed **4-5 page** case study document structured for engineering students.
    The content must be academic, rigorous, and strictly distinct from business-only marketing case studies.

    MANDATORY STRUCTURE (JSON):
    You must return a JSON object with the following exact structure. Do NOT include markdown formatting outside the JSON fields.
    
    {
      "title": "A Professional Title for the Case Study",
      "sections": [
        {
          "heading": "Introduction",
          "content": "Detailed background, context, and objectives..."
        },
        {
          "heading": "Case Description / Problem Statement",
          "content": "Clearly defined problem, stakeholder analysis, and evidence from the text..."
        },
        {
            "heading": "Analysis & Discussion",
            "content": "Stage 2 Rubric Core: Root cause analysis, application of relevant frameworks (SWOT/PESTEL/Technical Models), comparative evaluation of alternatives. Use critical thinking."
        },
        {
          "heading": "Findings & Observations",
          "content": "Key takeaways derived from the analysis..."
        },
        {
          "heading": "Conclusion & Recommendations",
          "content": "Actionable, feasible, and innovative recommendations linked to the problem statement."
        },
        {
          "heading": "SDG Impact Reflection",
          "content": "Identify relevant UN Sustainable Development Goals (SDGs). Explain expectations and contributions to social/economic/environmental impact."
        },
        {
          "heading": "Personal Reflection",
          "content": "A reflective section on learning outcomes, skills gained, and concept clarity."
        },
        {
          "heading": "References",
          "content": "List citations in IEEE or APA style."
        }
      ]
    }

    CONTENT RULES:
    1. **Strict Alignment**: Ensure the content satisfies Stage 1 (Problem Definition), Stage 2 (Critical Thinking & Frameworks), and Stage 3 (Clarity & Impact) of the standard case study rubric.
    2. **No Hallucinations**: Base all facts on the provided document. If information is missing, infer logically as an academic exercise but do not invent unrelated facts.
    3. **Academic Tone**: Use formal engineering/academic language.
    4. **Length**: The 'content' fields should be substantial, enabling the generation of a 4-5 page document.

    Return ONLY the valid JSON.
    """

    if mime_type.startswith("image/"):
        gemini_input_parts[1]["text"] += "\n\n" + prompt
    else:
        gemini_input_parts.append({"text": prompt})

    try:
        response = model.generate_content(gemini_input_parts)
        text_response = response.text
        
        # Clean markdown
        if text_response.startswith("```json"):
            text_response = text_response[7:].strip()
        if text_response.endswith("```"):
            text_response = text_response[:-3].strip()
            
        # Clean potential invalid control characters (newlines in strings are common issues)
        # Attempt to remove strictly invalid control characters while keeping newlines
        # Or, just use strict=False which handles some control chars
        
        # Simple fix: Escape unescaped newlines inside the JSON string is hard without a parser.
        # But valid JSON should have escaped newlines.
        # We will try strict=False first.
        
        return json.loads(text_response, strict=False)
    except json.JSONDecodeError as je:
        print(f"JSON Parse Error: {je}")
        print(f"Raw Response snippet: {text_response[:500]} ... {text_response[-500:]}")
        # Fallback: Try to sanitize common issues
        try:
             # Basic sanitation of control characters (0-31) except 9(tab), 10(nl), 13(cr)
             import re
             cleaned_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text_response)
             return json.loads(cleaned_text, strict=False)
        except Exception as e2:
             raise Exception(f"Gemini generation error (JSON): {je}. Cleaned failed too: {e2}")
    except Exception as e:
        raise Exception(f"Gemini generation error: {e}")

# --- DOCX Creation Logic ---
def create_case_study_docx(case_data, filename="sakec_case_study.docx"):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # --- Cover Page ---
    document.add_heading(case_data.get('title', 'Case Study'), level=0)
    document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\n" * 2)
    
    # Placeholders for Student Details
    details = [
        "Student Name: _______________________",
        "Roll No: ____________________________",
        "Department/Course: __________________",
        "Faculty Name: _______________________",
        "Institute: Shah & Anchor Kutchhi Engineering College (SAKEC)",
        "Date: _______________________________"
    ]
    for detail in details:
        p = document.add_paragraph(detail)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_page_break()

    # --- Table of Contents (Static Placeholder) ---
    document.add_heading("Table of Contents", level=1)
    for section in case_data.get('sections', []):
        document.add_paragraph(section.get('heading', 'Section'))
    document.add_page_break()

    # --- Content Sections ---
    for section in case_data.get('sections', []):
        heading = section.get('heading', '')
        content = section.get('content', '')
        
        if heading:
            document.add_heading(heading, level=1)
        
        if content:
            # Split content by newlines to preserve basic paragraphing if Gemini provided it
            paras = content.split('\n')
            for p_text in paras:
                if p_text.strip():
                    document.add_paragraph(p_text.strip())
            
        document.add_paragraph("\n") # Spacing between sections

    # Save
    file_path = os.path.join(CASE_STUDY_FOLDER, filename)
    document.save(file_path)
    return file_path

# --- Endpoint ---
@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "SAKEC Case Study Service Running", "port": 5016}), 200

@app.route("/generate-sakec-case-study", methods=["POST"])
def generate_sakec_case_study():
    if not request.json:
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")
    mime_type = request.json.get("mime_type")

    if not all([filename, base64_file_data, mime_type]):
        return jsonify({"error": "Missing fields"}), 400

    temp_path = None
    docx_path = None
    
    try:
        # Decode and Save Temp File
        file_bytes = base64.b64decode(base64_file_data)
        safe_name = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_path = os.path.join(UPLOAD_FOLDER, safe_name)
        
        with open(temp_path, "wb") as f:
            f.write(file_bytes)

        # Extract Text
        extracted_text = ""
        if mime_type.endswith('document') or filename.endswith('.docx'):
            extracted_text = extract_text_from_docx(temp_path)
        elif mime_type == 'application/pdf' or filename.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(temp_path)
        elif mime_type.startswith('image/'):
            extracted_text = extract_text_from_image(temp_path)
        
        # Generator Scenarios
        case_study_data = generate_sakec_case_study_content(extracted_text, mime_type, base64_file_data)
        
        # Create DOCX
        output_filename = f"SAKEC_Case_Study_{uuid.uuid4()}.docx"
        docx_path = create_case_study_docx(case_study_data, output_filename)
        
        # Return Result
        with open(docx_path, "rb") as f:
            b64_output = base64.b64encode(f.read()).decode('utf-8')
            
        return jsonify({
            "completed_filename": output_filename,
            "completed_file_data": b64_output
        })

    except Exception as e:
        print(f"ERROR: {e}")
        return jsonify({"error": str(e)}), 500
        
    finally:
        # Cleanup
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
        if docx_path and os.path.exists(docx_path):
            os.remove(docx_path)

if __name__ == "__main__":
    app.run(port=5016, debug=False)
